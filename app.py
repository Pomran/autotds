import streamlit as st
import pandas as pd
from docx import Document
import io
import datetime
import json
import os

# ==========================================
# 配置文件路径
# ==========================================
CONFIG_FILE = 'sample_gen_config.json'

# ==========================================
# 核心逻辑类 (保持 v7.0 逻辑不变)
# ==========================================

class SampleGenerator:
    def __init__(self):
        self.config = self.load_config()

    def get_default_config(self):
        return {
            "prefix_map": {
                "Transformer": "T", "Fuse": "B", "Resistor": "B", "Bleeder": "B",
                "X-Capacitor": "C", "X1": "C", "X2": "C",
                "Y-Capacitor": "D", "Y1": "D", "Y2": "D",
                "Opto": "E", "Photo": "E",
                "PCB": "F", "PWB": "F", "Board": "F",
                "Battery": "G",
                "Cable": "H", "Wire": "H", "Cord": "H", "Plug": "H", "Connector": "H",
                "Enclosure": "I", "Case": "I", "Plastic": "I", "Metal": "I",
                "Supply": "J", "Remote": "K", "Load": "L",
                "Inductor": "N", "Choke": "N", "Coil": "N", "Filter": "N",
                "Shade": "O",
                "Insulation Tape": "M", "Tape": "M", "Insulation": "M", 
                "Sheet": "M", "Silicone": "M", "Tube": "M"
            },
            "sort_order": ["T", "F", "B", "C", "D", "E", "N", "M", "I", "H", "G", "J", "K", "L", "Z"],
            "default_selection": ["Transformer", "Insulation"] 
        }

    def load_config(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return self.get_default_config()
        else:
            return self.get_default_config()

    def save_config(self, new_config):
        try:
            with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(new_config, f, indent=4, ensure_ascii=False)
            self.config = new_config
            return True
        except:
            return False

    def get_prefix(self, obj_name):
        current_map = self.config["prefix_map"]
        obj_upper = obj_name.upper()
        sorted_keys = sorted(current_map.keys(), key=len, reverse=True)
        for key in sorted_keys:
            if key.upper() in obj_upper:
                return current_map[key]
        return "Z"

    def clean_text(self, cell):
        if not cell: return ""
        text = cell.text.strip().replace('\n', ' ').replace('\r', '')
        text = text.lstrip("-").lstrip("•").strip()
        return text

    def parse_cdf_table(self, docx_file):
        doc = Document(docx_file)
        data = []
        target_table = None
        header_row_index = -1
        col_map = {} 

        # 智能寻找表头
        for table in doc.tables:
            if len(table.rows) > 0:
                for i, row in enumerate(table.rows[:6]): 
                    cells = row.cells
                    row_text_list = [c.text.lower().strip() for c in cells]
                    
                    has_object = any("object" in t or "part no" in t for t in row_text_list)
                    has_mfr = any("manufacturer" in t or "trademark" in t for t in row_text_list)
                    
                    if has_object and has_mfr:
                        target_table = table
                        header_row_index = i
                        for idx, text in enumerate(row_text_list):
                            if "object" in text or "part no" in text: col_map['object'] = idx
                            elif "manufacturer" in text or "trademark" in text: col_map['mfr'] = idx
                            elif "type" in text or "model" in text: col_map['model'] = idx
                            elif "technical" in text or "rating" in text or "data" in text: col_map['rating'] = idx
                            elif "standard" in text: col_map['standard'] = idx
                            elif "mark" in text or "conformity" in text: col_map['marks'] = idx
                        break
            if target_table: break
        
        if not target_table: return None, "未找到符合格式的元器件清单 (Table 1.A)"

        col_map.setdefault('object', 0)
        col_map.setdefault('mfr', 1)
        col_map.setdefault('model', 2)
        col_map.setdefault('rating', 3)
        col_map.setdefault('standard', 4)

        last_valid_object = ""
        
        for i, row in enumerate(target_table.rows[header_row_index + 1:]):
            try:
                cells = row.cells
                max_idx = max(v for k, v in col_map.items() if v is not None)
                if len(cells) <= max_idx: 
                    if 'marks' in col_map and col_map['marks'] >= len(cells): pass 
                    else: continue

                obj_name = self.clean_text(cells[col_map['object']])
                manufacturer = self.clean_text(cells[col_map['mfr']])
                raw_model_text = self.clean_text(cells[col_map['model']])
                tech_data = self.clean_text(cells[col_map['rating']])
                standard = self.clean_text(cells[col_map['standard']])
                marks = self.clean_text(cells[col_map['marks']]) if 'marks' in col_map and col_map['marks'] < len(cells) else ""

                if not manufacturer and raw_model_text:
                    upper_model = raw_model_text.upper()
                    company_keywords = ["CO LTD", "CO.,", "CORP", "INC.", "LIMITED", "GMBH", "S.R.L", "BV"]
                    if any(kw in upper_model for kw in company_keywords):
                        manufacturer = raw_model_text
                        raw_model_text = tech_data
                        tech_data = standard
                        standard = marks
                        marks = ""

                is_alternative = False
                if not obj_name: is_alternative = True
                elif "(Alternative)" in obj_name or "Alternative" in obj_name: is_alternative = True
                
                if is_alternative:
                    current_object = last_valid_object
                else:
                    current_object = obj_name
                    last_valid_object = obj_name 

                if not current_object: continue

                split_models = [m.strip() for m in raw_model_text.split(',') if m.strip()]
                if not split_models: split_models = [""]

                for single_model in split_models:
                    prefix_code = self.get_prefix(current_object)
                    
                    default_select = False
                    for kw in self.config["default_selection"]:
                        if kw.lower() in current_object.lower():
                            default_select = True
                            break

                    data.append({
                        "Select": default_select, 
                        "Prefix": prefix_code, 
                        "Object": current_object,
                        "Manufacturer": manufacturer,
                        "Model": single_model, 
                        "Ratings": tech_data,
                        "Standard": standard,
                        "Marks": marks
                    })

            except Exception: continue

        if not data: return None, "数据提取为空"
        return pd.DataFrame(data), None

    def generate_word_report(self, selected_df, project_no, product_info, date_received):
        doc = Document()
        date_str = date_received.strftime("%Y-%m-%d")
        
        heading = doc.add_heading(f'Test Sample List - Project {project_no}', 0)
        heading.alignment = 1 
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = False 
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sample Number'
        hdr_cells[1].text = 'Date Received'
        hdr_cells[2].text = 'Manufacturer, Product Identification and Ratings'
        for cell in hdr_cells: cell.width = 1500000 

        row = table.add_row().cells
        row[0].text = f"{project_no}-001~010"
        row[1].text = date_str  
        row[2].text = product_info

        counters = {} 
        for index, item in selected_df.iterrows():
            prefix = item['Prefix']
            if prefix not in counters: counters[prefix] = 1
            sample_id = f"{project_no}-001-{prefix}{counters[prefix]:03d}"
            
            raw_parts = [item['Object'], item['Manufacturer'], item['Model'], item['Ratings'], item['Standard'], item['Marks']]
            clean_parts = []
            seen = set()
            for part in raw_parts:
                if not part: continue
                p_clean = part.strip()
                if "Interchangeable" in p_clean: continue
                if p_clean not in seen:
                    clean_parts.append(p_clean)
                    seen.add(p_clean)
            
            full_desc = ", ".join(clean_parts)
            row = table.add_row().cells
            row[0].text = sample_id
            row[1].text = date_str 
            row[2].text = full_desc
            counters[prefix] += 1
            
        f_stream = io.BytesIO()
        doc.save(f_stream)
        f_stream.seek(0)
        return f_stream

def get_date_from_project_no(project_no):
    try:
        date_part = project_no[:6]
        if len(date_part) == 6 and date_part.isdigit():
            return datetime.datetime.strptime(date_part, "%y%m%d").date()
    except: pass
    return datetime.date.today()

# ==========================================
# Streamlit 界面
# ==========================================

def main():
    st.set_page_config(page_title="Sample List Generator", layout="wide")
    
    # 极简标题
    st.title("安规样品表生成工具")
    
    if 'generator' not in st.session_state:
        st.session_state.generator = SampleGenerator()
    generator = st.session_state.generator

    # --- Sidebar: 信息输入 ---
    with st.sidebar:
        st.subheader("项目信息")
        project_no = st.text_input("项目编号", value="251118001")
        default_date = get_date_from_project_no(project_no)
        date_received = st.date_input("收样日期", value=default_date)
        
        default_info = "POWERLD ENTERPRISES CO LTD\nPDF-150-15\nInput: 100-240V~, 50-60Hz, 2.0A\nOutput: 15VDC, 10A"
        product_info = st.text_area("整机规格描述", value=default_info, height=200)

    # --- Main: 标签页 ---
    tab_main, tab_config = st.tabs(["任务台", "参数配置"])

    # === Tab 2: 配置 (Config) ===
    with tab_config:
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**命名规则映射**")
            map_data = [{"Keyword": k, "Prefix": v} for k, v in generator.config["prefix_map"].items()]
            edited_map_df = st.data_editor(
                pd.DataFrame(map_data), 
                num_rows="dynamic", 
                column_config={
                    "Keyword": st.column_config.TextColumn("关键词", required=True),
                    "Prefix": st.column_config.TextColumn("字母", required=True, max_chars=2)
                },
                use_container_width=True,
                height=400,
                hide_index=True
            )

        with col2:
            st.markdown("**默认勾选**")
            sel_data = [{"Keyword": k} for k in generator.config["default_selection"]]
            edited_sel_df = st.data_editor(
                pd.DataFrame(sel_data),
                num_rows="dynamic",
                column_config={"Keyword": st.column_config.TextColumn("关键词", required=True)},
                use_container_width=True,
                height=150,
                hide_index=True
            )

            st.markdown("**排序优先级**")
            current_sort_str = ", ".join(generator.config["sort_order"])
            new_sort_str = st.text_area("输入字母顺序 (逗号分隔)", value=current_sort_str, height=100)

        st.write("")
        if st.button("保存配置", type="primary", use_container_width=True):
            new_config = {
                "prefix_map": dict(zip(edited_map_df['Keyword'], edited_map_df['Prefix'])),
                "sort_order": [s.strip().upper() for s in new_sort_str.split(',') if s.strip()],
                "default_selection": edited_sel_df['Keyword'].tolist()
            }
            if generator.save_config(new_config):
                st.success("配置已更新")
                st.experimental_rerun()

    # === Tab 1: 主任务 (Main) ===
    with tab_main:
        uploaded_file = st.file_uploader("上传元器件清单 (Table 1.A)", type=["docx"], label_visibility="collapsed")
        if not uploaded_file:
            st.caption("请上传 .docx 格式的 CDF 文件")

        if uploaded_file:
            with st.spinner('正在解析...'):
                df, error = generator.parse_cdf_table(uploaded_file)
            
            if error:
                st.error(error)
            else:
                # 排序逻辑
                sort_list = generator.config["sort_order"]
                def get_sort_index(prefix):
                    prefix = prefix.upper()
                    if prefix in sort_list: return sort_list.index(prefix)
                    return 999 + ord(prefix[0]) if prefix else 999

                df['SortIndex'] = df['Prefix'].apply(get_sort_index)
                df = df.sort_values(by=['SortIndex', 'Object'])

                # 编辑器
                st.divider()
                edited_df = st.data_editor(
                    df,
                    column_config={
                        "Select": st.column_config.CheckboxColumn("选择", width="small"),
                        "Prefix": st.column_config.TextColumn("分类", width="small"),
                        "Object": st.column_config.TextColumn("元件名称", width="medium"),
                        "Manufacturer": st.column_config.TextColumn("制造商", width="medium"),
                        "Model": st.column_config.TextColumn("型号", width="medium"),
                        "Ratings": st.column_config.TextColumn("规格", width="medium"),
                        "Marks": st.column_config.TextColumn("认证", width="small"),
                        "SortIndex": None
                    },
                    hide_index=True,
                    use_container_width=True,
                    height=500
                )

                selected_rows = edited_df[edited_df["Select"] == True]
                
                # 底部操作栏
                st.divider()
                c1, c2 = st.columns([3, 1])
                with c1:
                    if not selected_rows.empty:
                        st.caption(f"已选样品数: {len(selected_rows)}")
                    else:
                        st.caption("仅生成整机")
                
                with c2:
                    if st.button("导出 Word 文档", type="primary", use_container_width=True):
                        doc_stream = generator.generate_word_report(selected_rows, project_no, product_info, date_received)
                        st.download_button(
                            label="下载文件",
                            data=doc_stream,
                            file_name=f"{project_no}_Sample_List.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

if __name__ == "__main__":
    main()